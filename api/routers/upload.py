"""
api/routers/upload.py – 문서 업로드 및 파싱 라우터
POST /upload : 파일 업로드 → 텍스트 추출 → Gemini 필드 추출
POST /parse  : 직접 텍스트 입력 → Gemini 필드 추출
"""

import hashlib
import io
import logging
from typing import Annotated

import yaml
from fastapi import APIRouter, File, HTTPException, UploadFile, status
from fastapi.responses import JSONResponse
from pydantic import BaseModel

from api.config import get_settings
from api.services import gemini_service

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/upload", tags=["업로드 & 파싱"])


# ── 요청/응답 스키마 ────────────────────────────────────────────────────


class ParseRequest(BaseModel):
    """직접 텍스트 파싱 요청 바디"""
    text: str


class ParsedDocument(BaseModel):
    """파싱된 문서 응답"""
    doc_id: str
    filename: str | None = None
    raw_text: str
    fields: dict
    yaml_str: str
    char_count: int


# ── 문서 텍스트 추출 헬퍼 ──────────────────────────────────────────────


def _extract_text_from_pdf(content: bytes) -> str:
    """PDF 파일에서 텍스트를 추출합니다."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
        return "\n".join(texts)
    except Exception as exc:
        logger.error("PDF 텍스트 추출 오류: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"PDF 파싱 오류: {exc}",
        ) from exc


def _extract_text_from_docx(content: bytes) -> str:
    """DOCX 파일에서 텍스트를 추출합니다."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX 텍스트 추출 오류: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"DOCX 파싱 오류: {exc}",
        ) from exc


def _extract_text_from_xlsx(content: bytes) -> str:
    """XLSX 파일에서 텍스트를 추출합니다."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"[시트: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(c) if c is not None else "" for c in row)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as exc:
        logger.error("XLSX 텍스트 추출 오류: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"XLSX 파싱 오류: {exc}",
        ) from exc


def _extract_text(content: bytes, content_type: str, filename: str) -> str:
    """파일 형식에 따라 적합한 텍스트 추출기를 호출합니다."""
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        return _extract_text_from_pdf(content)
    elif (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    ):
        return _extract_text_from_docx(content)
    elif (
        content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        or filename.lower().endswith(".xlsx")
    ):
        return _extract_text_from_xlsx(content)
    elif content_type.startswith("text/"):
        return content.decode("utf-8", errors="replace")
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"지원하지 않는 파일 형식: {content_type}",
        )


# ── 엔드포인트 ───────────────────────────────────────────────────────


@router.post(
    "",
    response_model=ParsedDocument,
    summary="문서 업로드 및 AI 필드 추출",
    description="PDF/DOCX/XLSX 파일을 업로드하면 Gemini AI가 공시 필드를 자동 추출합니다.",
)
async def upload_document(
    file: Annotated[UploadFile, File(description="업로드할 문서 (PDF/DOCX/XLSX, 최대 50MB)")],
) -> ParsedDocument:
    """
    파일 업로드 → 텍스트 추출 → Gemini 필드 추출 → 응답 반환.

    에러 케이스:
    - 413: 파일 크기 초과
    - 415: 지원하지 않는 파일 형식
    - 422: 텍스트 추출 또는 AI 파싱 실패
    """
    settings = get_settings()

    # 파일 크기 검증
    content = await file.read()
    if len(content) > settings.max_upload_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"파일 크기({len(content):,}B)가 최대({settings.max_upload_bytes:,}B)를 초과합니다.",
        )

    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="파일이 비어 있습니다.",
        )

    # 파일 이름 안전 처리
    filename = file.filename or "unknown"
    content_type = file.content_type or ""

    logger.info("파일 업로드: '%s' (%d bytes, %s)", filename, len(content), content_type)

    # 텍스트 추출
    raw_text = _extract_text(content, content_type, filename)

    if not raw_text.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="문서에서 텍스트를 추출할 수 없습니다. (이미지 기반 PDF 또는 암호화된 파일)",
        )

    # Gemini AI 필드 추출
    try:
        fields = await gemini_service.extract_fields(raw_text)
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(exc),
        ) from exc
    except RuntimeError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=str(exc),
        ) from exc

    # YAML 직렬화
    yaml_str = yaml.dump(fields, allow_unicode=True, default_flow_style=False)

    # 문서 ID (파일 내용 해시 기반)
    doc_id = hashlib.sha256(content).hexdigest()[:16]

    return ParsedDocument(
        doc_id=doc_id,
        filename=filename,
        raw_text=raw_text[:2000],  # 응답 크기 제한
        fields=fields,
        yaml_str=yaml_str,
        char_count=len(raw_text),
    )


@router.post(
    "/parse",
    response_model=ParsedDocument,
    summary="텍스트 직접 파싱",
    description="직접 입력한 텍스트에서 Gemini AI가 공시 필드를 추출합니다.",
)
async def parse_text(body: ParseRequest) -> ParsedDocument:
    """
    텍스트 입력 → Gemini 필드 추출 → 응답 반환.
    """
    text = body.text.strip()
    if not text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="텍스트가 비어 있습니다.",
        )

    try:
        fields = await gemini_service.extract_fields(text)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    yaml_str = yaml.dump(fields, allow_unicode=True, default_flow_style=False)
    doc_id = hashlib.sha256(text.encode()).hexdigest()[:16]

    return ParsedDocument(
        doc_id=doc_id,
        filename=None,
        raw_text=text[:2000],
        fields=fields,
        yaml_str=yaml_str,
        char_count=len(text),
    )
